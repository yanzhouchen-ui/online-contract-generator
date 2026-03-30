"""
Word 文档处理工具函数
从原项目提取，适配线上环境
"""

from docx import Document
from docx.shared import Pt


def format_date_spanish(date_obj):
    """将日期格式化为西班牙语格式"""
    months = {
        1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
        5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
        9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
    }
    return f"{date_obj.day} de {months[date_obj.month]} de {date_obj.year}"


def apply_style(run):
    """应用格式: Times New Roman, Size 10"""
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)


def replace_text_in_doc(doc, replacements):
    """
    替换 Word 文档中的文本
    
    Args:
        doc: Document 对象
        replacements: 字典，key 是占位符（如 [Name]），value 是替换值
    """
    # 替换段落中的文本
    for paragraph in doc.paragraphs:
        # 优化: 快速跳过不含方括号的段落
        if "[" not in paragraph.text:
            continue
        for key, value in replacements.items():
            if key in paragraph.text:
                # 策略 1: 尝试在 run 级别替换，以保留部分格式
                replaced_in_run = False
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, str(value))
                        apply_style(run)
                        replaced_in_run = True
                
                # 策略 2: 如果 run 级别没找到 (说明 key 被打散了)，则进行整段替换
                if not replaced_in_run:
                    paragraph.text = paragraph.text.replace(key, str(value))
                    # 整段替换后，重新应用格式给所有 run
                    for run in paragraph.runs:
                        apply_style(run)

    # 替换表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # 优化: 快速跳过不含方括号的段落
                    if "[" not in paragraph.text:
                        continue
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            # 同样的策略
                            replaced_in_run = False
                            for run in paragraph.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, str(value))
                                    apply_style(run)
                                    replaced_in_run = True
                            
                            if not replaced_in_run:
                                paragraph.text = paragraph.text.replace(key, str(value))
                                for run in paragraph.runs:
                                    apply_style(run)
                        # 尝试去除软回车/换行符后匹配 (针对字段被换行打断的情况)
                        elif key in paragraph.text.replace('\n', '').replace('\r', ''):
                            clean_text = paragraph.text.replace('\n', '').replace('\r', '')
                            paragraph.text = clean_text.replace(key, str(value))
                            for run in paragraph.runs:
                                apply_style(run)

    # 替换页眉和页脚
    for section in doc.sections:
        # 处理页眉
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for paragraph in header.paragraphs:
                    # 优化: 快速跳过
                    if "[" not in paragraph.text:
                        continue
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(value))
                            for run in paragraph.runs:
                                apply_style(run)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                # 优化: 快速跳过
                                if "[" not in paragraph.text:
                                    continue
                                for key, value in replacements.items():
                                    if key in paragraph.text:
                                        paragraph.text = paragraph.text.replace(key, str(value))
                                        for run in paragraph.runs:
                                            apply_style(run)
        
        # 处理页脚
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for paragraph in footer.paragraphs:
                    # 优化: 快速跳过
                    if "[" not in paragraph.text:
                        continue
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(value))
                            for run in paragraph.runs:
                                apply_style(run)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                # 优化: 快速跳过
                                if "[" not in paragraph.text:
                                    continue
                                for key, value in replacements.items():
                                    if key in paragraph.text:
                                        paragraph.text = paragraph.text.replace(key, str(value))
                                        for run in paragraph.runs:
                                            apply_style(run)
